"""
Actualiza Informe_Avances_Rediseñado.docx con los resultados CAL-15
(run 2026-05-01: outputs/run_20260501_cal15_full.log).

Convierte la versión CAL-8 (Apr 2026) en versión CAL-9 → CAL-10b.2 → CAL-12 →
CAL-13 → CAL-14 → CAL-15. Preserva formato (negritas, fuentes, estilos) y solo
sustituye texto donde hay datos numéricos o frases ahora desactualizadas.

Convención de seguridad:
- En tablas: para cada celda objetivo se reescribe `cell.text = new_str`,
  python-docx mantiene el estilo del primer run.
- En párrafos: se hace `replace_in_runs(p, old, new)` que itera runs y aplica
  reemplazo solo si el patrón completo cabe en un run; si está fragmentado,
  hace un fallback prudente que une todo el texto en el primer run y limpia
  los demás (preserva la fuente del primer run).
"""
from __future__ import annotations
import sys, io, shutil
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")

from copy import deepcopy
import docx
from docx.oxml.ns import qn

PATH_IN = "Informe_Avances_Rediseñado.docx"
PATH_BAK = "Informe_Avances_Rediseñado_pre_cal15.docx"


def replace_in_runs(paragraph, old: str, new: str) -> bool:
    """
    Sustituye `old` → `new` en los runs del párrafo. Devuelve True si hubo
    cambio. Estrategia:
      1. Buscar `old` dentro de un único run (caso simple).
      2. Si no se encuentra dentro de un run pero sí en el texto unido,
         consolidar texto en el primer run, limpiar los demás.
    """
    # Caso simple: cabe en un solo run
    for r in paragraph.runs:
        if old in r.text:
            r.text = r.text.replace(old, new)
            return True
    # Fallback: texto fragmentado
    full = paragraph.text
    if old not in full:
        return False
    new_full = full.replace(old, new)
    if not paragraph.runs:
        paragraph.text = new_full
        return True
    # Conservar el primer run + limpiar el resto.
    paragraph.runs[0].text = new_full
    for r in paragraph.runs[1:]:
        r.text = ""
    return True


def set_cell_text(cell, new_text: str) -> None:
    """Reescribe el texto de la celda preservando el estilo del primer run."""
    if not cell.paragraphs:
        cell.text = new_text
        return
    p0 = cell.paragraphs[0]
    # Si la celda tiene múltiples párrafos, vaciar los siguientes.
    for p in cell.paragraphs[1:]:
        p._element.getparent().remove(p._element)
    if not p0.runs:
        p0.text = new_text
        return
    p0.runs[0].text = new_text
    for r in p0.runs[1:]:
        r.text = ""


def update_table_row(table, row_idx: int, cells: list[str | None]):
    """
    Actualiza una fila de tabla. Cada elemento de `cells` puede ser:
      - str: nuevo texto para esa celda
      - None: dejar sin cambios
    """
    row = table.rows[row_idx]
    for c_idx, new_text in enumerate(cells):
        if new_text is None:
            continue
        if c_idx >= len(row.cells):
            continue
        set_cell_text(row.cells[c_idx], new_text)


def main():
    shutil.copy(PATH_IN, PATH_BAK)
    print(f"[backup] {PATH_BAK}")

    d = docx.Document(PATH_IN)

    # ════════════════════════════════════════════════════════════════════
    # 1) PORTADA: subtítulo CAL-8 → CAL-15
    # ════════════════════════════════════════════════════════════════════
    # P013: "Recalibración con tarifas Cedenar mensuales per-agente — CAL-8"
    target = d.paragraphs[13]
    replace_in_runs(target,
                     "Recalibración con tarifas Cedenar mensuales per-agente — CAL-8",
                     "Liquidación regulatoria post-CAL-15: C4 hereda CREG 174 art. 25 (Tipo 1/Tipo 2 + Cvm)")

    # T00: AUTOR/ASESORES/FECHA  (mes)
    update_table_row(d.tables[0], 0, [
        None,  # AUTOR/Brayan
        None,  # ASESORES
        "FECHA | Mayo de 2026 | San Juan de Pasto",
    ])

    # ════════════════════════════════════════════════════════════════════
    # 2) RESUMEN EJECUTIVO — Síntesis cuantitativa
    # ════════════════════════════════════════════════════════════════════
    # P037 — actualizar mención CAL-8 → CAL-9..CAL-15 + horizonte
    p37 = d.paragraphs[37]
    replace_in_runs(p37,
                     "Como novedad principal de esta versión, los precios de venta al usuario provienen de las tarifas reales mensuales de Cedenar (calibración CAL-8, 13 PDFs respaldatorios desde abril 2025 hasta abril 2026).",
                     "Como novedad principal de esta versión, los escenarios C1 y C4 implementan la dualidad regulatoria CREG 174/2021 art. 25 (CAL-10b.2 y CAL-15): la permuta intracomunitaria se valora a (π_gs − Cvm,i,j) y los excedentes a precio de bolsa horario. C2 está alineado con la Ley 143/1994 + CREG 086/1996 + Decreto 388/2007 (CAL-13), y el precio de bolsa lleva techo CREG 101 066/2024 PES (CAL-14). La tarifa π_gs es matriz mensual per-agente (CAL-9, 13 PDFs Cedenar abril 2025 – abril 2026).")

    # T02 — caja de calibración CAL-8
    t02 = d.tables[2]
    set_cell_text(t02.rows[0].cells[0],
                   "CALIBRACIÓN CAL-15 — QUÉ CAMBIÓ | El escenario C4 (AGRC) hereda formalmente CREG 174/2021 art. 25 vía Decreto 2236/2023 art. 4: la permuta intracomunitaria Tipo 1 se valora a (π_gs − Cvm) y el excedente residual Tipo 2 a precio de bolsa horario. Esto cierra la asimetría entre C1 (que ya implementaba el descuento Cvm desde CAL-10b.2) y C4 (que pre-CAL-15 valoraba créditos PDE a π_gs completo). Numéricamente C4 sube +4,03 % y queda casi empatado con P2P; la ventaja P2P pasa de +2,14 MCOP (RPE +4,08 %) a +132 K (RPE +0,25 %).")

    # P040 — bullet inicial: agregar π_bolsa con techo PES
    # (no cambia esencialmente, lo dejamos)

    # T03 — Caja KPI 4 columnas
    update_table_row(d.tables[3], 0, [
        "GANANCIA P2P | 52,4M | COP por período",
        "GANANCIA C4 ★ | 52,3M | Baseline regulatorio (CAL-15)",
        "RPE P2P VS C4 | +0,25 % | Empate técnico",
        "VENTAJA ABSOLUTA | +132 K | COP P2P − C4",
    ])

    # T04 — Tabla 1: comparación regulatoria global
    # Encabezado intacto. Filas R1..R5: ganancias post-CAL-15.
    update_table_row(d.tables[4], 1, [  # R1: C1
        "C1 — CREG 174/2021",
        "52 603 335",
        "0,176",
        "0,921",
        "+0,0136",
    ])
    update_table_row(d.tables[4], 2, [  # R2: P2P
        "P2P (Stackelberg + RD)",
        "52 446 938",
        "0,188",
        "0,981",
        "+0,3677",
    ])
    update_table_row(d.tables[4], 3, [  # R3: C2
        "C2 — PPA bilateral",
        "51 437 446",
        "0,176",
        "0,921",
        "+0,0296",
    ])
    update_table_row(d.tables[4], 4, [  # R4: C3
        "C3 — Mercado spot",
        "50 958 336",
        "0,176",
        "0,921",
        "+0,0379",
    ])
    update_table_row(d.tables[4], 5, [  # R5: C4
        "C4 — CREG 101 072 ★",
        "52 314 652",
        "0,176",
        "0,921",
        "+0,0568",
    ])

    # P044 — caption Tabla 1
    p44 = d.paragraphs[44]
    replace_in_runs(p44,
                     "Tabla 1. Comparación regulatoria global post-CAL-8 (calibración Cedenar mensual per-agente). SC = self-consumption, SS = self-sufficiency, IE = índice de equidad. C1 emerge como dominante absoluto pero P2P captura el segundo lugar con ventaja distributiva clara.",
                     "Tabla 1. Comparación regulatoria global post-CAL-15. SC = self-consumption, SS = self-sufficiency, IE = índice de equidad. C1 sigue siendo el de mayor monto absoluto pero P2P y C4 quedan estadísticamente empatados (Δ = +132 K, RPE = +0,25 %); la ventaja P2P se desplaza al Nivel 2 (equidad: IE_P2P = +0,3677 vs IE_C4 = +0,0568).")

    # P046, P047, P048 — bullets clave
    p46 = d.paragraphs[46]
    replace_in_runs(p46,
                     "El P2P supera a C4 en +2 140 790 COP (RPE = +4,08 %), confirmando que la asignación dinámica captura beneficio que el PDE estático no puede asignar. La ventaja absoluta casi duplicó la del análisis pre-CAL-8 gracias a la heterogeneidad oficial/comercial.",
                     "El P2P supera a C4 por apenas +132 286 COP (RPE = +0,25 %). Tras CAL-15 el escenario AGRC reconoce permuta intracomunitaria a (π_gs − Cvm) hora a hora y exporta el residual a π_bolsa[k]; esto sube C4 en +4,03 % vs el modelo pre-CAL-15 (50,29 → 52,31 MCOP) y cierra la brecha cuantitativa. La superioridad P2P se sostiene principalmente en la dimensión distributiva.")
    p47 = d.paragraphs[47]
    replace_in_runs(p47,
                     "C1 emerge como el escenario de mayor monto absoluto (54,04 MCOP, +1,61 sobre P2P) porque la permutación 1:1 ahora se valoriza al π_gs[n] específico de cada agente — mucho mayor que el escalar 650 anterior. Su PoF = 0,000 (eficiente y equitativo coinciden).",
                     "C1 sigue siendo el escenario de mayor monto absoluto (52,60 MCOP, +156 K sobre P2P) tras CAL-10b.2. La permuta Tipo 1 se valoriza a (π_gs[n] − Cvm,i,j) por linealidad CREG 174 art. 25; el excedente Tipo 2 va a π_bolsa[k] horario post-Hx. Su PoF = 0,000 (eficiente y equitativo coinciden — ambos en C1).")
    p48 = d.paragraphs[48]
    replace_in_runs(p48,
                     "Hallazgo regulatorio nuevo: Udenar y HUDN (los dos agentes oficiales) entran en zona de deserción individual frente a C1. La frontera relevante para la racionalidad individual ya no es C4 sino C1 — se discute con detalle en §8.4.",
                     "Hallazgo regulatorio post-CAL-15: la racionalidad individual queda fragmentada en tres frentes — Udenar prefiere C1 (+1,16 MCOP sobre P2P); Mariana, UCC y HUDN prefieren C4 (entre +61 K y +164 K sobre P2P); solo Cesmag se mantiene estable en P2P (+137 K sobre la mejor alternativa regulatoria). Detalle en §8.4.")

    # P050 — caption Figura 1
    p50 = d.paragraphs[50]
    replace_in_runs(p50,
                     "Figura 1. Comparación de la ganancia neta agregada por escenario regulatorio post-CAL-8. C1 alcanza el mayor monto absoluto (54,04 MCOP) pero el P2P captura el segundo lugar con ventaja distributiva clara (IE = +0,3677), distribuyendo el excedente 71,4 % a compradores y 28,6 % a vendedores.",
                     "Figura 1. Comparación de la ganancia neta agregada por escenario regulatorio post-CAL-15. C1 lidera (52,60 MCOP) seguido por P2P (52,45) y C4 (52,31) en empate técnico; C2 y C3 quedan ~1,4 MCOP por debajo. La distintiva del P2P es ahora el componente distributivo (IE = +0,3677, distribución 72,0 / 28,0 % comprador/vendedor).")

    # ════════════════════════════════════════════════════════════════════
    # 3) Sección 1.1 — Perfiles MTE (no cambian las medias salvo cobertura)
    # ════════════════════════════════════════════════════════════════════
    # T05 — Tabla 2: caracterización empírica (sin cambios numéricos
    # estructurales; cobertura agregada se preserva)

    # ════════════════════════════════════════════════════════════════════
    # 4) Sección 1.2 — Precios y parámetros base
    # ════════════════════════════════════════════════════════════════════
    # T06 — Tabla 3: tarifas oficial/comercial. Mantenemos las medias 792
    # y 950 que son el promedio simple de los CU mensuales del CSV.
    # Solo aclarar que esas medias son aproximadas para diagnóstico — el
    # settlement usa la matriz (N,T) mes a mes.

    # T07 — Tabla 4: parámetros monetarios y temporales. Actualizar:
    # R3 π_gb=280 → seguimos con 280 nominal, pero π_bolsa real medio = 216.2
    # R4 π_bolsa medio: 280 → 216
    # R7 horizonte: ya correcto (6144)
    update_table_row(d.tables[7], 4, [   # R4 fila π_bolsa medio
        "Precio bolsa medio",
        "π̄_bolsa",
        "216 (post-PES)",
        "COP/kWh",
        "XM Colombia + techo CREG 101 066/2024 (CAL-14)",
    ])
    update_table_row(d.tables[7], 5, [   # R5 PPA — CAL-13
        "Precio PPA C2",
        "π_ppa",
        "≈ 401",
        "COP/kWh",
        "Bilateral usuario no-regulado (CAL-13, Ley 143/1994 + CREG 086/1996)",
    ])

    # ════════════════════════════════════════════════════════════════════
    # 5) Sección 3 — C1 (CREG 174)
    # ════════════════════════════════════════════════════════════════════
    # T10 — Tabla 5: métricas C1
    update_table_row(d.tables[10], 1, [   # R1 ganancia C1
        "Ganancia neta C1 agregada",
        "52 603 335 COP",
        "Mayor entre los 5 escenarios",
    ])
    update_table_row(d.tables[10], 2, [   # R2 sobre P2P
        "Sobre P2P",
        "+156 397 COP",
        "C1 vence P2P por margen estrecho (CAL-10b.2)",
    ])
    update_table_row(d.tables[10], 3, [   # R3 sobre C4
        "Diferencia con C4",
        "+288 683 COP",
        "C1 sigue por encima de C4 tras CAL-15",
    ])
    update_table_row(d.tables[10], 4, [   # R4 IE
        "IE (índice de equidad)",
        "+0,0136",
        "Casi neutral, leve sesgo a compradores",
    ])

    # T11 — Tabla 6: per-agente C1 vs C4 con valores CAL-15
    # Pre-CAL-15: C1=10,536,467 etc. Post-CAL-15: C1=9,277,305 etc; C4 nuevo
    update_table_row(d.tables[11], 1, [   # Udenar
        "Udenar (oficial)",
        "9 277 305",
        "7 784 811",
        "+1 492 494",
    ])
    update_table_row(d.tables[11], 2, [   # Mariana
        "Mariana (comercial)",
        "11 980 573",
        "12 362 441",
        "−381 868",
    ])
    update_table_row(d.tables[11], 3, [   # UCC
        "UCC (comercial)",
        "14 677 856",
        "15 281 043",
        "−603 187",
    ])
    update_table_row(d.tables[11], 4, [   # HUDN
        "HUDN (oficial)",
        "10 230 562",
        "10 375 707",
        "−145 145",
    ])
    update_table_row(d.tables[11], 5, [   # Cesmag
        "Cesmag (comercial)",
        "6 437 039",
        "6 510 651",
        "−73 612",
    ])
    update_table_row(d.tables[11], 6, [   # Total
        "Total comunidad",
        "52 603 335",
        "52 314 652",
        "+288 683",
    ])

    # P108 — caption Tabla 6
    p108 = d.paragraphs[108]
    replace_in_runs(p108,
                     "Tabla 6. Ganancia por institución bajo C1 vs C4 (post-CAL-8). Udenar muestra la mayor diferencia absoluta porque su π_gs oficial (792 COP/kWh) y su perfil con muchas horas excedentarias hacen que la permutación mensual capitalice mucho más valor que el PDE estático.",
                     "Tabla 6. Ganancia por institución bajo C1 vs C4 (post-CAL-15). Solo Udenar mantiene ventaja en C1; las cuatro instituciones restantes obtienen mayor beneficio bajo C4. Tras CAL-15 el AGRC reconoce permuta intracomunitaria con descuento Cvm,i,j igual que el AGPE, lo que iguala numéricamente ambos regímenes excepto en Udenar (mayor exportador con π_gs oficial).")

    # ════════════════════════════════════════════════════════════════════
    # 6) Sección 4 — C2 (PPA, CAL-13)
    # ════════════════════════════════════════════════════════════════════
    # P116 — π_ppa mid-point: ahora bajo CAL-13 es bilateral no-regulado
    p116 = d.paragraphs[116]
    replace_in_runs(p116,
                     "En el modelo se asume π_ppa como el punto medio entre π_gb y π̄_gs ponderado, es decir 593 COP/kWh para los parámetros del estudio post-CAL-8. Esta elección refleja la práctica de los mercados PPA donde el comprador busca un descuento frente a la tarifa regulada y el vendedor busca una prima frente al precio mayorista.",
                     "Tras CAL-13 (alineación con Ley 143/1994 art. 11 + CREG 086/1996 art. 1 + Decreto 388/2007 art. 4), C2 modela un contrato bilateral de un usuario no-regulado agregado: la comunidad MTE supera el umbral de 0,1 MW de potencia conjunta y 55 MWh-mes de consumo agregado, por lo que es elegible al régimen no-regulado. El precio efectivo π_ppa ≈ 401 COP/kWh resulta del punto medio (π_gb + π̄_gs)/2 con π_gs valorado solo en su componente G + Cvm + COT (no se duplican peajes T+D+PR+Rm que el comercializador-respaldo sigue cobrando aunque la energía la suministre un tercero).")

    # P124 — resultado C2
    p124 = d.paragraphs[124]
    replace_in_runs(p124,
                     "Con los datos MTE post-CAL-8, la ganancia agregada del PPA es 51 440 813 COP. Este valor está ligeramente por debajo del P2P (52,43 MCOP) pero por encima de C3 y C4. El comportamiento estructural del PPA se mantiene: el resultado depende casi enteramente del volumen intercambiado, no del precio pactado, porque cualquier aumento de π_ppa beneficia al vendedor a costa del comprador y viceversa. La diferencia entre PPA y P2P no está en cuánto valor se genera, sino en cómo se distribuye internamente.",
                     "Con los datos MTE post-CAL-13, la ganancia agregada del PPA es 51 437 446 COP. Este valor está por debajo del P2P (52,45 MCOP), C1 (52,60) y C4 (52,31), por encima de C3 (50,96). El comportamiento estructural del PPA se mantiene: el resultado depende casi enteramente del volumen intercambiado, no del precio pactado, porque cualquier aumento de π_ppa beneficia al vendedor a costa del comprador y viceversa. La diferencia entre PPA y P2P no está en cuánto valor se genera, sino en cómo se distribuye internamente.")

    # T12 — descomposición C2 (R2, R4 cambian valores)
    update_table_row(d.tables[12], 2, [   # R2 ingreso PPA
        "Ingreso PPA",
        "Prosumidor",
        "π_ppa ≈ 401",
    ])
    update_table_row(d.tables[12], 4, [   # R4 excedente residual
        "Excedente residual",
        "Prosumidor",
        "π_bolsa[k] horario (post-PES, CAL-14)",
    ])

    # ════════════════════════════════════════════════════════════════════
    # 7) Sección 5 — C3 (spot)
    # ════════════════════════════════════════════════════════════════════
    p137 = d.paragraphs[137]
    replace_in_runs(p137,
                     "Con los datos MTE y precios bolsa calibrados, la ganancia agregada de C3 es 50 961 703 COP, ligeramente por debajo de C2 y por encima de C4. La razón estructural se mantiene: la cobertura PV es del 19 % y los pocos excedentes que se generan en horas solares se venden a un precio bolsa que en esas horas suele ser bajo (la generación solar nacional deprime los precios diurnos).",
                     "Con los datos MTE y precios bolsa calibrados (post-CAL-14, techo PES aplicado a 12 horas / 0,23 % del horizonte), la ganancia agregada de C3 es 50 958 336 COP, por debajo de los cuatro escenarios restantes. La razón estructural se mantiene: la cobertura PV es del 19 % y los pocos excedentes que se generan en horas solares se venden a un precio bolsa que en esas horas suele ser bajo (la generación solar nacional deprime los precios diurnos).")

    # ════════════════════════════════════════════════════════════════════
    # 8) Sección 6 — C4 (CREG 101 072 / Decreto 2236)
    # ════════════════════════════════════════════════════════════════════
    # P144 — importancia para la tesis
    p144 = d.paragraphs[144]
    replace_in_runs(p144,
                     "Este es el escenario regulatorio vigente en Colombia para autogeneración colectiva (AGRC) y, por tanto, la referencia obligada contra la cual debe contrastarse el mercado P2P. Su marco normativo es el Decreto 2236 de 2023 y la Resolución CREG 101 072 de 2025.",
                     "Este es el escenario regulatorio vigente en Colombia para autogeneración colectiva (AGRC) y, por tanto, la referencia obligada contra la cual debe contrastarse el mercado P2P. Su marco normativo es el Decreto 2236 de 2023 art. 4 y la Resolución CREG 101 072 de 2025 art. 5, los cuales establecen que cada miembro AGRC se liquida bajo el régimen de Generador Distribuido y AGPE — es decir, hereda CREG 174/2021. Por linealidad regulatoria (CAL-15, 2026-05-01), C4 distingue Excedentes Tipo 1 (permuta intracomunitaria, valorada a π_gs − Cvm,i,j) y Tipo 2 (residual, a π_bolsa[k] horario), idéntico criterio que C1.")

    # P146 — ya describe correctamente PDE estático; no tocamos.

    # T15 — métodos PDE: ya correctos.

    # P155 — Resultados C4 (sección 6.4) — el siguiente párrafo a 6.4 es T16
    # T16 — métricas C4
    update_table_row(d.tables[16], 1, [
        "Spread ineficiencia estática",
        "1 004,4 kWh/período",
        "Energía mal asignada por PDE",
    ])
    update_table_row(d.tables[16], 2, [
        "RPE (P2P vs C4)",
        "+0,25 %",
        "Empate técnico tras CAL-15",
    ])
    update_table_row(d.tables[16], 3, [
        "IE (C4)",
        "+0,0568",
        "Casi neutral, leve sesgo a compradores",
    ])
    update_table_row(d.tables[16], 4, [
        "Score robustez C4",
        "1,00",
        "Cumple plenamente CREG 101 072",
    ])

    # ════════════════════════════════════════════════════════════════════
    # 9) Sección 7 — Comparación cruzada
    # ════════════════════════════════════════════════════════════════════
    # T17 — vista comparativa: actualizar fila C4
    update_table_row(d.tables[17], 4, [
        "C4 — CREG 101 072",
        "Hora a hora vía PDE + cruce comercial",
        "(π_gs − Cvm) en T1 + π_bolsa[k] en T2 (CAL-15)",
        "PDE estático sub-óptimo frente a descalce horario",
    ])

    # T18 — IE por escenario
    update_table_row(d.tables[18], 1, [   # P2P
        "P2P", "+0,3677", "Pro-comprador fuerte", "72,0 % a compradores",
    ])
    update_table_row(d.tables[18], 2, [   # C4
        "C4", "+0,0568", "Casi neutral", "Leve sesgo a compradores",
    ])
    update_table_row(d.tables[18], 3, [   # C3
        "C3", "+0,0379", "Casi neutral", "Leve sesgo a compradores",
    ])
    update_table_row(d.tables[18], 4, [   # C2
        "C2", "+0,0296", "Casi neutral", "Reparto cuasi simétrico",
    ])
    update_table_row(d.tables[18], 5, [   # C1
        "C1", "+0,0136", "Casi neutral", "Permuta T1 a (π_gs−Cvm) — sesgo leve a compradores",
    ])

    # P172 — PoF coincidencia
    p172 = d.paragraphs[172]
    replace_in_runs(p172,
                     "Sobre la comunidad MTE post-CAL-8, el escenario eficiente y el equitativo coinciden ambos en C1 (mayor monto absoluto y menor coeficiente de Gini), por lo que PoF = 0,000: no hay tensión entre eficiencia y equidad sobre este horizonte.",
                     "Sobre la comunidad MTE post-CAL-15, el escenario eficiente y el equitativo coinciden ambos en C1 (mayor monto absoluto 52,60 MCOP y menor coeficiente de Gini 0,1459), por lo que PoF = 0,000: no hay tensión entre eficiencia y equidad sobre este horizonte. C4 tiene el Gini más alto (0,1691) entre los cinco — la distribución administrativa por PDE concentra ligeramente más el beneficio en los grandes generadores.")

    # T19 — P2P vs C4 per agente — CRÍTICO: 4 signos invierten
    update_table_row(d.tables[19], 1, [   # Udenar — sigue P2P > C4
        "Udenar (oficial)",
        "8 117 765",
        "7 784 811",
        "+332 954",
    ])
    update_table_row(d.tables[19], 2, [   # Mariana — INVIERTE
        "Mariana (comercial)",
        "12 198 377",
        "12 362 441",
        "−164 064",
    ])
    update_table_row(d.tables[19], 3, [   # UCC — INVIERTE
        "UCC (comercial)",
        "15 219 747",
        "15 281 043",
        "−61 295",
    ])
    update_table_row(d.tables[19], 4, [   # HUDN — INVIERTE
        "HUDN (oficial)",
        "10 263 249",
        "10 375 707",
        "−112 458",
    ])
    update_table_row(d.tables[19], 5, [   # Cesmag — sigue P2P > C4
        "Cesmag (comercial)",
        "6 647 801",
        "6 510 651",
        "+137 150",
    ])
    update_table_row(d.tables[19], 6, [   # Total
        "Total comunidad",
        "52 446 938",
        "52 314 652",
        "+132 286",
    ])

    # P176 — texto crítico "5/5 prefieren P2P"
    p176 = d.paragraphs[176]
    replace_in_runs(p176,
                     "La siguiente tabla muestra cómo cada institución se comporta bajo el escenario P2P y bajo C4 (la referencia regulatoria vigente). Las cinco prefieren P2P sobre C4 (racionalidad individual 5/5 frente a C4):",
                     "La siguiente tabla muestra cómo cada institución se comporta bajo el escenario P2P y bajo C4 (la referencia regulatoria vigente post-CAL-15). Solo 2/5 prefieren P2P sobre C4 (Udenar y Cesmag); las otras tres (Mariana, UCC, HUDN) obtienen mayor beneficio bajo C4 — el mecanismo AGRC con permuta intracomunitaria reconocida supera al P2P en estas tres por margen estrecho (61 K – 164 K COP):")

    # P177 — caption Tabla 12
    p177 = d.paragraphs[177]
    replace_in_runs(p177,
                     "Tabla 12. Ventaja P2P vs C4 por institución (post-CAL-8). UCC, el agente con mayor demanda y tarifa comercial, captura la mayor ventaja absoluta.",
                     "Tabla 12. Ventaja P2P vs C4 por institución (post-CAL-15). Tras heredar CREG 174 art. 25 a través del Decreto 2236, C4 captura el ahorro de Mariana, UCC y HUDN; Udenar y Cesmag mantienen ventaja P2P por sus perfiles asimétricos de demanda. Total agregado: P2P > C4 por +132 K COP (RPE +0,25 %).")

    # ════════════════════════════════════════════════════════════════════
    # 10) Sección 8 — Sensibilidad
    # ════════════════════════════════════════════════════════════════════
    # T20 — SA-1: nueva tabla CAL-15
    rows_sa1 = [
        ("200", "53 367 398", "53 024 813", "+0,006"),
        ("250", "53 184 432", "53 145 898", "+0,001"),
        ("280", "53 074 653", "53 218 549", "−0,003"),
        ("300", "53 001 466", "53 266 982", "−0,005"),
        ("350", "52 818 501", "53 388 067", "−0,011"),
        ("400", "52 635 535", "53 509 152", "−0,017"),
        ("450", "52 452 569", "53 630 236", "−0,022"),
        ("500", "52 269 604", "53 751 321", "−0,028"),
    ]
    for i, vals in enumerate(rows_sa1, start=1):
        update_table_row(d.tables[20], i, list(vals))

    # P194 — caption SA-1
    p194 = d.paragraphs[194]
    replace_in_runs(p194,
                     "Tabla 13. Sensibilidad del rendimiento P2P vs C4 al precio de bolsa π_gb (post-CAL-8). Fila resaltada: caso base. C4 es invariante al π_bolsa porque con cobertura PV agregada de 11,3 % nunca hay excedente comunitario para liquidar a bolsa.",
                     "Tabla 13. Sensibilidad del rendimiento P2P vs C4 al precio de bolsa π_gb (post-CAL-15). Tras heredar Tipo 2 a π_bolsa horario, C4 ahora SÍ varía con π_bolsa (a diferencia del modelo pre-CAL-15). El RPE pasa de positivo (PGB ≤ 250) a negativo (PGB ≥ 280): a precios de bolsa altos C4 captura más Tipo 2 y supera al P2P estructuralmente.")

    # T22 — Tabla deserción individual: 5 filas, valores CAL-15
    update_table_row(d.tables[22], 1, [   # Udenar
        "Udenar", "oficial",
        "8 118 k", "9 277 k (C1)", "−12,5 %", "180",
    ])
    update_table_row(d.tables[22], 2, [   # Mariana — ahora C4 alt
        "Mariana", "comercial",
        "12 198 k", "12 362 k (C4)", "−1,3 %", "180",
    ])
    update_table_row(d.tables[22], 3, [   # UCC — ahora C4 alt
        "UCC", "comercial",
        "15 220 k", "15 281 k (C4)", "−0,4 %", "180",
    ])
    update_table_row(d.tables[22], 4, [   # HUDN — sigue C1 alt? log dice C4
        "HUDN", "oficial",
        "10 263 k", "10 376 k (C4)", "−1,1 %", "180",
    ])
    update_table_row(d.tables[22], 5, [   # Cesmag — único estable
        "Cesmag", "comercial",
        "6 648 k", "6 511 k (C4)", "+2,1 %", "494",
    ])

    # T21 — caja "frontera de racionalidad cambia"
    set_cell_text(d.tables[21].rows[0].cells[0],
                   "FRONTERA DE RACIONALIDAD INDIVIDUAL POST-CAL-15 | Tras CAL-15, la comunidad como un todo apenas prefiere P2P sobre C4 (+132 K COP, RPE = +0,25 %). A nivel individual, solo Cesmag se mantiene estable en P2P. Udenar prefiere C1 (−12,5 % bajo P2P respecto a 9,28 MCOP en C1), y Mariana, UCC, HUDN prefieren C4 (entre −0,4 % y −1,3 % bajo P2P respecto al beneficio C4 reconocido por la permuta intracomunitaria CREG 174 art. 25). La frontera de deserción agregada se desplaza por debajo de π_gb = 280 COP/kWh: para 4/5 agentes la mejor alternativa regulatoria ya no es C4-pre-CAL-15 sino C1-post-CAL-10b.2 o C4-post-CAL-15 con descuento Cvm + permuta intracomunitaria.")

    # P204 — caption Tabla 14
    p204 = d.paragraphs[204]
    replace_in_runs(p204,
                     "Tabla 14. Análisis de racionalidad individual post-CAL-8. Los umbrales críticos π_gb* de Udenar (180) y HUDN (233) están dentro del rango histórico de bolsa XM 2025. 3/5 agentes son estables en P2P, 2/5 (oficiales) prefieren C1.",
                     "Tabla 14. Análisis de racionalidad individual post-CAL-15. Solo Cesmag (1/5) es estable en P2P. Udenar prefiere C1 con margen amplio (−12,5 %); Mariana, UCC y HUDN prefieren C4 con margen estrecho (−0,4 % a −1,3 %). El umbral π_gb* = 180 COP/kWh para 4 agentes está dentro del rango histórico XM 2025; el de Cesmag (494) sería un escenario de bolsa anormalmente alta.")

    # P208 — explicación económica
    p208 = d.paragraphs[208]
    replace_in_runs(p208,
                     "La explicación analítica del hallazgo es directa: en C1, cada kWh excedente permutado se valoriza al π_gs[n] específico del agente. Para los agentes oficiales (792 COP/kWh), eso es lo que reciben directamente. En P2P el vendedor obtiene una prima (π* − π_gb) que es estructuralmente menor que su propio π_gs[n]. En cambio, los agentes comerciales (950 COP/kWh) tienen un π_gs[n] alto que les permite extraer más ahorro como compradores en P2P (π_gs[i] − π*[i]), manteniendo Δ_n positivo.",
                     "La explicación analítica del hallazgo es directa: tras CAL-15, C4 valoriza la permuta intracomunitaria a (π_gs[n] − Cvm[n]) en lugar de π_gs completo, pero compensa esa caída unitaria con un volumen mayor de permuta reconocida (la inyección bruta a la frontera comunitaria, no el neto agregado). Para Mariana, UCC y HUDN —cuyos perfiles tienen poca generación propia y mucha demanda compatible con permuta intracomunitaria— C4 captura más valor que P2P. Solo Udenar mantiene C1 como mejor alternativa porque su perfil tiene muchas horas excedentarias absolutas que la permuta C1 mensual capitaliza completamente. Cesmag, con cobertura PV equilibrada y demanda contenida, es el único agente donde la flexibilidad horaria del P2P supera ambos regímenes administrativos.")

    # T23 — implicación regulatoria
    set_cell_text(d.tables[23].rows[0].cells[0],
                   "IMPLICACIÓN REGULATORIA POST-CAL-15 | El P2P ya no domina automáticamente al AGRC en monto agregado: con CAL-15 (CREG 174 art. 25 heredada vía Decreto 2236) C4 reconoce permuta intracomunitaria con descuento Cvm,i,j y queda en empate técnico con P2P. La aportación distintiva del P2P se desplaza a tres frentes: (i) Nivel 2 — equidad fuerte (IE = +0,3677 vs +0,0568 de C4); (ii) flexibilidad regulatoria (FA-3: si un participante se retira, el P2P preserva flexibility premium positivo en 4/5 casos); (iii) valor de la asignación dinámica horaria — el P2P responde al descalce hora a hora, mientras el PDE estático no puede premiar a quien efectivamente necesita la energía en una hora específica. Para defensa de tesis: la narrativa cambia de 'P2P vence a C4 por +4 % cuantitativo' a 'P2P empata cuantitativamente con C4 pero captura el excedente comunitario distributivamente'. Este matiz fortalece la fidelidad regulatoria del modelo.")

    # ════════════════════════════════════════════════════════════════════
    # 11) Sección 9 — Conclusiones
    # ════════════════════════════════════════════════════════════════════
    p212 = d.paragraphs[212]
    replace_in_runs(p212,
                     "El mercado P2P (Stackelberg + RD) supera de manera consistente al escenario regulatorio vigente C4 (CREG 101 072/2025) sobre los datos empíricos MTE post-CAL-8. La ventaja absoluta es +2,14 MCOP/período (RPE = +4,08 %), casi el doble del análisis pre-CAL-8 gracias a la heterogeneidad oficial/comercial. Las cinco instituciones prefieren P2P sobre C4.",
                     "El mercado P2P (Stackelberg + RD) y el escenario regulatorio vigente C4 (CREG 101 072/2025) quedan estadísticamente empatados sobre los datos empíricos MTE post-CAL-15. La ventaja absoluta es apenas +132 K COP/período (RPE = +0,25 %), muy por debajo del +2,14 MCOP reportado pre-CAL-15. Solo 2/5 instituciones prefieren P2P sobre C4 a nivel individual: Udenar y Cesmag. Mariana, UCC y HUDN prefieren el AGRC tras la corrección regulatoria CAL-15.")

    p213 = d.paragraphs[213]
    replace_in_runs(p213,
                     "C1 (CREG 174/2021) emerge como el escenario de mayor monto absoluto (54,04 MCOP) porque la permutación 1:1 ahora se valoriza al π_gs[n] específico de cada agente: 792 oficial / 950 comercial. Su IE = −0,0115 es prácticamente neutral, y su PoF = 0,000 indica que sobre la comunidad MTE no hay tensión entre eficiencia y equidad.",
                     "C1 (CREG 174/2021) emerge como el escenario de mayor monto absoluto (52,60 MCOP) tras CAL-10b.2 (descuento Cvm,i,j literal sobre la permuta Tipo 1, art. 25). Su IE = +0,0136 es prácticamente neutral, y su PoF = 0,000 indica que sobre la comunidad MTE no hay tensión entre eficiencia y equidad — el más equitativo (menor Gini) y el más eficiente coinciden en C1.")

    p214 = d.paragraphs[214]
    replace_in_runs(p214,
                     "C2 (PPA) y C3 (spot) producen valores intermedios (51,44 vs 50,96 MCOP). El PPA es zero-sum interno: el resultado depende casi enteramente del volumen intercambiado, no del precio pactado. C3 está limitado por la correlación negativa entre disponibilidad solar y precio bolsa diurno.",
                     "C2 (PPA, ahora bilateral no-regulado bajo CAL-13) y C3 (spot) producen valores 51,44 y 50,96 MCOP respectivamente. El PPA es zero-sum interno: el resultado depende casi enteramente del volumen intercambiado, no del precio pactado. C3 está limitado por la correlación negativa entre disponibilidad solar y precio bolsa diurno; el techo PES (CREG 101 066/2024 vía CAL-14) recorta 12 horas de pico (0,23 % del horizonte).")

    p215 = d.paragraphs[215]
    replace_in_runs(p215,
                     "C4 (CREG 101 072) presenta la peor performance absoluta (50,29 MCOP) porque el PDE estático no puede premiar a quien efectivamente necesita la energía en una hora específica. El spread de ineficiencia estática es de 1 004,4 kWh/período.",
                     "C4 (CREG 101 072) tras CAL-15 sube a 52,31 MCOP (+4,03 % vs el 50,29 pre-CAL-15) al reconocer permuta intracomunitaria con descuento Cvm,i,j y excedente residual a π_bolsa[k] horario. El spread de ineficiencia estática se mantiene en 1 004,4 kWh/período: el PDE sigue siendo estático y no puede premiar a quien efectivamente necesita la energía en una hora específica.")

    p216 = d.paragraphs[216]
    replace_in_runs(p216,
                     "Hallazgo regulatorio nuevo post-CAL-8: Udenar y HUDN (los dos agentes oficiales) tendrían incentivo a desertar del P2P hacia C1. La frontera relevante para la racionalidad individual ya no es C4 sino C1, lo cual indica que el diseño regulatorio del P2P en Colombia debería incluir un mecanismo compensatorio para agentes oficiales.",
                     "Hallazgo regulatorio post-CAL-15: la frontera de racionalidad individual queda fragmentada — Udenar prefiere C1 (gran exportador con perfil oficial), Mariana/UCC/HUDN prefieren C4 (alta demanda, permuta intracomunitaria los favorece tras CAL-15), Cesmag se mantiene estable en P2P. El diseño regulatorio del P2P en Colombia debe incluir mecanismos compensatorios diferenciados según el perfil estructural del agente — una talla única no resuelve la fragmentación.")

    p217 = d.paragraphs[217]
    replace_in_runs(p217,
                     "La ventaja del P2P escala con la cobertura PV. Conforme la comunidad MTE incremente capacidad instalada, el spread de ineficiencia estática del PDE crecerá y el RPE seguirá favoreciendo al P2P. Este resultado es robusto al precio de bolsa y al subperíodo evaluado.",
                     "La ventaja distributiva del P2P (IE +0,3677 vs +0,0568 de C4) escala con la cobertura PV: a mayor potencia instalada, el descalce horario se intensifica y el P2P captura más excedente comunitario hora a hora. Este resultado es robusto al precio de bolsa: el RPE pasa de levemente positivo a levemente negativo en SA-1, pero el IE y la robustez FA-3/FA-4 se sostienen en todo el rango.")

    p218 = d.paragraphs[218]
    replace_in_runs(p218,
                     "La calibración con tarifas Cedenar reales (CAL-8) amplificó cuantitativamente la mayoría de los efectos respecto al pre-CAL-8: el bienestar absoluto subió ~37 % en los cinco escenarios y la ventaja absoluta P2P-C4 casi se duplicó, pero las jerarquías cualitativas se mantuvieron (C1 ≥ P2P > C2 ≥ C3 > C4).",
                     "Las correcciones regulatorias acumuladas CAL-9 (π_gs matriz N×T) → CAL-10b.2 (Cvm,i,j literal en C1) → CAL-12/13 (C2 Front-of-Meter no-regulado) → CAL-14 (techo PES) → CAL-15 (Tipo 1/2 + Cvm en C4) produjeron tres cambios estructurales sobre el reporte pre-CAL-8: (i) los montos absolutos cayeron en C1 (54,04 → 52,60 MCOP) por el descuento Cvm; (ii) C4 subió (50,29 → 52,31 MCOP) por la permuta intracomunitaria; (iii) C2 bajó marginalmente (51,44 → 51,44, igual a primera vista pero con descomposición regulatoria distinta). La nueva jerarquía es C1 > P2P ≈ C4 > C2 > C3, con el P2P diferenciándose por equidad y flexibilidad más que por monto absoluto.")

    # Footer del informe
    p260 = d.paragraphs[260]
    replace_in_runs(p260,
                     "Tesis P2P · Universidad de Nariño · Abril 2026",
                     "Tesis P2P · Universidad de Nariño · Mayo 2026 · Post-CAL-15")

    # ════════════════════════════════════════════════════════════════════
    # GUARDAR
    # ════════════════════════════════════════════════════════════════════
    d.save(PATH_IN)
    print(f"[saved] {PATH_IN}")
    print(f"[backup pre-CAL-15] {PATH_BAK}")


if __name__ == "__main__":
    main()
