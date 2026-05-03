"""
Dump del Informe_Avances_Rediseñado.docx con todo su contenido a UTF-8.
Útil para localizar los párrafos/tablas con números numéricos a actualizar.
"""
from __future__ import annotations
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")

import docx

d = docx.Document("Informe_Avances_Rediseñado.docx")

# Mapa: para cada elemento del cuerpo, indicar tipo y posición.
# python-docx no expone document order entre paragraphs y tables fácilmente,
# pero podemos iterar element.body y reconstruir el orden.
from docx.oxml.ns import qn

para_idx = -1
table_idx = -1

for child in d.element.body.iterchildren():
    if child.tag == qn("w:p"):
        para_idx += 1
        # Encontrar el paragraph correspondiente
        p = d.paragraphs[para_idx]
        style = p.style.name if p.style else "None"
        text = p.text.strip()
        if not text and not (style and style.startswith("Heading")):
            continue
        marker = ""
        if style and style.startswith("Heading"):
            marker = f" [{style}]"
        print(f"\n--- P{para_idx:03d}{marker} ---")
        print(text)
    elif child.tag == qn("w:tbl"):
        table_idx += 1
        if table_idx >= len(d.tables):
            continue
        t = d.tables[table_idx]
        print(f"\n--- T{table_idx:02d} ({len(t.rows)} filas × "
              f"{len(t.rows[0].cells) if t.rows else 0} cols) ---")
        for r_idx, row in enumerate(t.rows):
            cells = [c.text.strip().replace("\n", " | ") for c in row.cells]
            print(f"  R{r_idx}: " + " │ ".join(cells))
